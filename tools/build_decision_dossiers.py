# -*- coding: utf-8 -*-
"""Um dossiê de decisão por perfil e por ano, baixável do site.

O produto promete que cada recomendação vira um documento que se defende
sozinho anos depois. Esta rotina cumpre a promessa para as onze decisões
retrospectivas de cada perfil declarado: um documento por decisão, com o que
era conhecido na data, a informação posterior deliberadamente separada, os
parâmetros congelados da política e o hash do registro que impede a reescrita.

As fontes são as mesmas do site — ``web/composition.json``, ``web/ladder_v2.json``
e o registro congelado — para que nenhum número do PDF possa divergir da página
que o oferece. A saída são .docx em ``artifacts/decision_dossiers/docx``; a
conversão a PDF (LibreOffice headless) publica em ``web/dossiers``.
"""
from __future__ import annotations

from pathlib import Path
import json
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from portfolio_risk import risk_profile_spec  # noqa: E402

OUT_DOCX = ROOT / "artifacts" / "decision_dossiers" / "docx"
NAVY = RGBColor(0x10, 0x2A, 0x43)
TEAL = RGBColor(0x0C, 0x80, 0x76)
MUTED = RGBColor(0x5B, 0x6F, 0x7C)
LABELS = {"conservador": "Conservador", "equilibrado": "Equilibrado", "arrojado": "Arrojado"}
ACTIONS = {"entered": "entrou", "maintained": "mantido", "increased": "aumentou",
           "reduced": "reduziu", "exited": "saiu", "not_held": "fora"}


def pct(value, digits=2):
    if value is None:
        return "—"
    return f"{value * 100:.{digits}f}%".replace(".", ",")


def shade(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def para(document, text, *, size=10, bold=False, color=None, space_after=6, align=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return paragraph


def heading(document, text):
    return para(document, text, size=13, bold=True, color=NAVY, space_after=4)


def kicker(document, text):
    return para(document, text.upper(), size=8, bold=True, color=TEAL, space_after=2)


def add_table(document, headers, rows, widths=None, right_from=1):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header.upper())
        run.font.size = Pt(7.5)
        run.bold = True
        run.font.color.rgb = MUTED
        shade(cell, "F1F5F3")
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            if j >= right_from:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8.5)
            if j == 0:
                run.bold = True
                run.font.color.rgb = NAVY
    if widths:
        for j, width in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Cm(width)
    return table


def annual_result(curve: dict, series: str, year: int) -> float | None:
    dates, values = curve["dates"], curve["series"][series]
    def level_at_end(y):
        candidates = [v for d, v in zip(dates, values) if int(d[:4]) == y]
        return candidates[-1] if candidates else None
    end = level_at_end(year)
    start = level_at_end(year - 1)
    if end is None:
        return None
    base = start if start is not None else 100.0
    return end / base - 1


def build_dossier(profile: str, block: dict, ladder: dict, registration: dict) -> Document:
    label = LABELS[profile]
    year = int(block["decision_year"])
    date_br = "/".join(reversed(block["decision_date"].split("-")))
    declared = ladder["profiles"][profile]["declared"]
    spec = risk_profile_spec(profile)
    overlay = registration["intrayear_overlay"]["config"]
    sha = registration["registration_sha256"]

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.8)
    section.left_margin = section.right_margin = Cm(2.0)

    kicker(document, f"Benevente · política {registration['policy']} · documento de demonstração")
    para(document, "Dossiê de decisão de carteira", size=20, bold=True, color=NAVY, space_after=2)
    para(document, f"Perfil {label} · decisão de {date_br}", size=13, color=TEAL, space_after=10)

    status = document.add_table(rows=1, cols=1)
    status.style = "Table Grid"
    cell = status.rows[0].cells[0]
    shade(cell, "E8F6F2")
    cell.text = ""
    run = cell.paragraphs[0].add_run("RECONSTRUÇÃO RETROSPECTIVA SOB POLÍTICA CONGELADA\n")
    run.font.size = Pt(8); run.bold = True; run.font.color.rgb = TEAL
    run = cell.paragraphs[0].add_run(
        f"A política que produz esta decisão foi declarada e congelada em 26/08/2026 "
        f"(registro SHA-256 {sha[:16]}…, aprovado por Diego Gallina); as decisões de 2015 a 2025 são a "
        f"reconstrução do que ela teria feito, usando apenas informação disponível em cada data. "
        f"Nenhum parâmetro muda depois do congelamento e a amostra confirmatória começa no primeiro "
        f"pregão de 2027. Este documento é demonstração de pesquisa, não recomendação de investimento.")
    run.font.size = Pt(8.5)

    para(document, "", size=4, space_after=2)
    heading(document, "1 · Parâmetros declarados do perfil")
    add_table(document,
              ["Parâmetro", "Valor declarado"],
              [["Orçamento de renda variável", pct(declared["maximum_equity_weight"], 0)],
               ["Emissores na cesta", str(declared["top_assets"])],
               ["Teto por emissor", pct(declared["maximum_asset_weight"], 1)],
               ["Perna global (IVVB11, S&P 500 em reais)", f"{pct(declared['global_share_of_portfolio'], 0)} do patrimônio (um quinto do orçamento de ações)"],
               ["Teto setorial", "3 emissores por setor CVM"],
               ["Sinal de seleção", "Qualidade, valor e momento de doze meses (fatores pré-declarados)"],
               ["Revisão da cesta", "Anual, no primeiro pregão de janeiro"],
               ["Camada de proteção (só perna doméstica)",
                f"alerta: queda {pct(overlay['alert_drawdown'], 0)} ou vol. {pct(overlay['alert_volatility'], 0)} → expõe {pct(spec.alert_multiplier, 0)} do alvo; "
                f"severo: {pct(overlay['severe_drawdown'], 0)}/{pct(overlay['severe_volatility'], 0)} → {pct(spec.severe_multiplier, 0)}; "
                f"retomada após {overlay['recovery_days']} pregões calmos"]],
              widths=[7.2, 9.4])

    para(document, "", size=4, space_after=2)
    heading(document, f"2 · Alocação da decisão de {year}")
    add_table(document,
              ["Bloco", "Fração do patrimônio"],
              [["Ações Brasil (cesta selecionada)", pct(block["domestic_equity"], 0)],
               ["S&P 500 em reais (IVVB11, declarado)", pct(block["global_sleeve"], 0)],
               ["Caixa remunerado (Tesouro Selic)", pct(block["cash"], 0)]],
              widths=[9.0, 7.6])
    para(document, "A camada de proteção pode reduzir a fração doméstica dentro do ano, conforme os "
                   "gatilhos da Seção 1; a perna global e o caixa não participam do sinal.",
         size=8.5, color=MUTED)

    heading(document, "3 · Posições — o que era conhecido na data")
    rows = []
    for position in block["positions"]:
        rows.append([position["ticker"].removesuffix(".SA"),
                     pct(position["weight"]),
                     "—" if position.get("previous_weight") is None else pct(position["previous_weight"]),
                     ACTIONS.get(position["action"], position["action"]),
                     "—" if position.get("score") is None else f"{position['score']:.3f}".replace(".", ","),
                     pct(position.get("trailing_12m"), 1),
                     pct(position.get("trailing_vol"), 1)])
    add_table(document, ["Emissor", "Peso", "Peso anterior", "Ação", "Score", "Retorno 12m", "Vol. 12m"],
              rows, widths=[2.6, 2.3, 2.6, 2.3, 2.1, 2.4, 2.3])
    para(document, "Score, retorno e volatilidade de doze meses eram observáveis na data da decisão. "
                   "Peso anterior compara com o janeiro anterior do mesmo perfil.",
         size=8.5, color=MUTED)

    document.add_page_break()
    kicker(document, "Informação posterior · separada de propósito")
    heading(document, "4 · O que aconteceu depois")
    para(document, "Nada nesta seção estava disponível na data da decisão, e nada dela pode "
                   "retroalimentar a regra. Ela existe para auditoria do resultado, não para a decisão.",
         size=9, color=MUTED)
    result = annual_result(ladder["monthly_curve"], LABELS[profile], year)
    # A política vigente declara o caixa como instrumento; o dossiê compara
    # contra o papel que o cliente compra, com o nome dele.
    cash_label = next(k for k in ladder["monthly_curve"]["series"]
                      if k not in (*LABELS.values(), "Ibovespa"))
    cdi = annual_result(ladder["monthly_curve"], cash_label, year)
    ibov = annual_result(ladder["monthly_curve"], "Ibovespa", year)
    global_row = block.get("global_row") or {}
    add_table(document, ["Série", f"Resultado do ciclo {year}"],
              [[f"Perfil {label} (com proteção)", pct(result)],
               ["Perna global (IVVB11)", pct(global_row.get("realised_next_year"))],
               [cash_label, pct(cdi)],
               ["Ibovespa (retorno total)", pct(ibov)]],
              widths=[9.0, 7.6])
    rows = [[position["ticker"].removesuffix(".SA"), pct(position.get("realised_next_year"), 1)]
            for position in block["positions"]]
    add_table(document, ["Emissor", "Retorno no ciclo seguinte"], rows, widths=[5.0, 6.0])

    heading(document, "5 · Custos e imposto")
    para(document,
         "Os resultados publicados são líquidos de custos modelados: taxas de negociação, deslizamento "
         "proporcional à participação no volume médio diário e o custo das trocas da camada de proteção. "
         "O imposto de renda variável é apurado por ativo e custo médio, mês a mês, com isenção de "
         "R$ 20 mil para ações, compensação de prejuízos e ETF sem isenção (rotina "
         "tax_lot_accounting; arrasto medido de 0,7 a 2,1 pontos ao ano conforme o perfil). "
         "Falta a conciliação com nota de corretagem, que exige operação real.", size=9)

    heading(document, "6 · Governança")
    para(document,
         "O sistema calcula e explica; não transmite ordens. Em operação real, cada dossiê carrega o hash "
         "dos insumos, a versão da política e a aprovação nominal de um profissional habilitado — a decisão "
         "é de quem assina. O modelo de linguagem apenas transforma fatos já aprovados em tese e riscos "
         "para revisão humana; desligá-lo não altera nenhum peso.", size=9)

    heading(document, "7 · Limitações")
    for item in (
        "A janela 2015–2025 também orientou fatores e restrições: é amostra de desenvolvimento, não teste prospectivo.",
        "A camada de proteção foi desenhada depois das crises da amostra e reage com uma sessão de atraso.",
        "Cerca de um terço do retorno da perna global veio da desvalorização do real; é posição em dólar sem proteção cambial.",
        "Parte das séries históricas depende de proventos imputados; a reconciliação primária integral está pendente.",
        "Retornos retrospectivos não constituem recomendação, promessa de rentabilidade nem evidência de superioridade futura.",
    ):
        paragraph = para(document, f"•  {item}", size=8.5, color=MUTED, space_after=3)
        paragraph.paragraph_format.left_indent = Cm(0.3)

    para(document, "", size=4)
    kicker(document, "Verificação")
    para(document,
         f"Registro da política: SHA-256 {sha} · aprovado por Diego Gallina (diegogallina1) em 25/08/2026 · "
         f"fonte dos dados desta página: web/composition.json e web/ladder_v2.json, os mesmos arquivos do site "
         f"benevente.dgo.fi · gerado por tools/build_decision_dossiers.py em 26/08/2026.",
         size=7.5, color=MUTED)
    return document


def main() -> None:
    composition = json.loads((ROOT / "web" / "composition.json").read_text(encoding="utf-8"))
    ladder = json.loads((ROOT / "web" / "ladder_v2.json").read_text(encoding="utf-8"))
    registration = json.loads((ROOT / "data" / "benevente_profile_ladder_v3_registration.json").read_text(encoding="utf-8"))
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    count = 0
    for profile, blocks in composition["profiles"].items():
        for block in blocks:
            document = build_dossier(profile, block, ladder, registration)
            name = f"dossie_{profile}_{block['decision_year']}.docx"
            document.save(OUT_DOCX / name)
            count += 1
    print(f"{count} dossiês em {OUT_DOCX}")


if __name__ == "__main__":
    main()
