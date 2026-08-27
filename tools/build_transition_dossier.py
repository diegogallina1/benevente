# -*- coding: utf-8 -*-
"""O dossiê do plano: o documento que justifica ter mexido — ou não mexido.

O dossiê de decisão anual explica por que uma carteira é o que é. Este explica
outra coisa, e mais delicada: por que valeu a pena sair de onde o cliente estava,
ou por que valeu a pena ficar. São perguntas diferentes, e a segunda é a que
volta anos depois, porque mexer custa dinheiro na hora e o benefício, se existir,
aparece devagar.

Por isso o documento tem duas seções que o outro não tem.

A primeira é o **caminho não escolhido**. Um documento que só descreve a opção
tomada não registra uma decisão, registra uma execução: quem lê depois não tem
como saber se houve escolha. Os dois caminhos aparecem lado a lado, com custo e
alcance de cada um, e o escolhido é marcado.

A segunda é a **assinatura**. A decisão anual é a política rodando sozinha; a
mudança é alguém decidindo aplicá-la a um cliente específico, num dia
específico, sabendo o que custa. Um dossiê de plano sem nome e sem data é
exatamente o papel que não serve para nada quando alguém pergunta, três anos
depois, quem mandou vender.

E o custo vem antes do resto, na primeira página, porque a ordem inversa é como
se vende uma troca ruim.
"""
from __future__ import annotations

from pathlib import Path
import argparse
import hashlib
import json
import sys

from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from tools.build_decision_dossiers import (MUTED, NAVY, TEAL, add_table, heading,  # noqa: E402
                                           kicker, para, shade)

OUT_DOCX = ROOT / "artifacts" / "transition_dossiers"
ACOES = {"vender": "Vender", "comprar": "Comprar", "reduzir": "Reduzir", "manter": "Manter"}
CESTAS = {"renda_variavel": "Renda variável", "renda_fixa": "Renda fixa",
          "fora_do_escopo": "Fora do escopo"}


def brl(value: float) -> str:
    return f"R$ {value:,.0f}".replace(",", ".")


def pct(value: float, casas: int = 1) -> str:
    return f"{value * 100:.{casas}f}%".replace(".", ",")


def _caixa(document: Document, titulo: str, texto: str, cor: str) -> None:
    tabela = document.add_table(rows=1, cols=1)
    tabela.style = "Table Grid"
    cell = tabela.rows[0].cells[0]
    shade(cell, cor)
    cell.text = ""
    run = cell.paragraphs[0].add_run(titulo + "\n")
    run.font.size = Pt(8); run.bold = True; run.font.color.rgb = NAVY
    run = cell.paragraphs[0].add_run(texto)
    run.font.size = Pt(8.5)


def build(payload: dict, cliente: str, escolha: str, registro: dict) -> Document:
    escolhido = payload["mapping"] if escolha == "adequar" else payload["alternative"]
    outro = payload["alternative"] if escolha == "adequar" else payload["mapping"]
    perfil = payload["intake"]["assessment"]
    nome_perfil = perfil["profile"]

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.7)
    section.left_margin = section.right_margin = Cm(2.0)

    kicker(document, f"Benevente · política {registro['policy']} · documento de demonstração")
    para(document, "Dossiê do plano", size=20, bold=True, color=NAVY, space_after=2)
    para(document, f"{cliente} · perfil {nome_perfil} · carteira de {brl(escolhido['total_brl'])}",
         size=13, color=TEAL, space_after=8)
    para(document, f"Plano escolhido: {escolhido['path_label']}.", size=10, bold=True,
         space_after=8)

    completo = escolhido.get("tax_is_complete", True)
    _caixa(document,
           "O QUE ISSO CUSTA, E É PAGO AGORA" if completo
           else "O QUE ISSO CUSTA — PISO, NÃO TOTAL",
           ("" if completo else "A partir de ")
           + f"{brl(escolhido['transition_total_brl'])}, ou "
           f"{pct(escolhido['transition_cost_pct'], 2)} do patrimônio: "
           f"{brl(escolhido['transition_cost_brl'])} de execução e "
           f"{brl(escolhido['transition_tax_brl'])} de imposto que a venda realiza. Uma vez, no "
           f"momento da mudança. Este documento não estima em quanto tempo isso se recupera — "
           f"fazê-lo exigiria projetar retorno futuro.", "FBF2E2")

    if not completo:
        para(document, "", size=4, space_after=2)
        _caixa(document, "O IMPOSTO ACIMA ESTÁ INCOMPLETO",
               f"{brl(escolhido['unpriced_sale_brl'])} estão sendo vendidos sem custo de "
               f"aquisição apurável: "
               + "; ".join(f"{p['ticker']} ({p['cost_quality']})"
                           for p in escolhido["positions_without_cost_basis"])
               + ". A B3 não fornece preço médio, e a base dela começa em 01/11/2019: posição "
                 "mais antiga chega sem as compras que a formaram. O valor não foi estimado de "
                 "propósito — um imposto estimado tem a mesma aparência de um medido e leva à "
                 "mesma decisão de vender. A nota de corretagem antiga fecha a conta.", "F6E9E4")

    if not escolhido["track_record_applies"]:
        para(document, "", size=4, space_after=2)
        _caixa(document, "O HISTÓRICO PUBLICADO NÃO DESCREVE ESTA CARTEIRA",
               "Os retornos que a Benevente publica foram medidos com seleção e proteção juntas. "
               "Este plano mantém a seleção do cliente e aplica só a proteção. Não existe "
               "medição do que a proteção sozinha teria feito sobre uma cesta escolhida por "
               "terceiro, e nenhum número deste projeto pode ser lido como previsão do resultado "
               "desta carteira.", "F6E9E4")

    para(document, "", size=4, space_after=2)
    heading(document, "1 · As respostas que definiram o perfil")
    para(document, "Não há pontuação: cada resposta impõe um teto de perfil e vale o menor deles. "
                   "Por isso sempre dá para apontar qual resposta determinou o resultado.", size=9)
    limites = perfil["all_limits"]
    if limites:
        add_table(document, ["Pergunta", "Resposta", "Teto imposto", "Por quê"],
                  [[l["question"], l["answer"], l["caps_at"], l["why"] or "—"] for l in limites],
                  widths=[2.6, 4.4, 2.6, 7.0])
    para(document, f"{perfil['rationale']} A pior queda medida deste perfil na janela declarada "
                   f"foi de {pct(perfil['worst_measured_drawdown'])}.", size=9)

    heading(document, "2 · A carteira que chegou")
    para(document, f"{brl(escolhido['total_brl'])} no total, "
                   f"{pct(payload['mapping']['alignment'])} já de acordo com o que a política "
                   f"declara para o perfil {nome_perfil}.", size=9.5)
    add_table(document, ["Origem dos dados"], [[o] for o in escolhido["sources"]],
              widths=[16.6], right_from=9)

    heading(document, "3 · O que muda")
    linhas = [[ACOES.get(m["action"], m["action"]), m["ticker"], brl(m["from_brl"]),
               brl(m["to_brl"]), m["reason"]]
              for m in escolhido["moves"] if m["action"] != "manter"]
    if linhas:
        add_table(document, ["Ação", "Ativo", "De", "Para", "Motivo declarado"], linhas,
                  widths=[2.0, 2.6, 3.0, 3.0, 6.0])
    else:
        para(document, "Nada muda: a carteira já está dentro de todos os limites deste plano.",
             size=9)

    mantidos = [m for m in escolhido["moves"] if m["action"] == "manter"]
    if mantidos:
        para(document, "", size=4, space_after=2)
        para(document, "Permanecem sem alteração: "
             + "; ".join(f"{m['ticker']} ({brl(m['from_brl'])}, {m['reason']})" for m in mantidos)
             + ".", size=8.5, color=MUTED)

    heading(document, "4 · O que custa, em detalhe")
    custo = [["Execução (corretagem, emolumentos e deslizamento)",
              brl(escolhido["transition_cost_brl"])]]
    for cesta, dados in escolhido["tax_by_bucket"].items():
        nome = CESTAS.get(cesta, cesta)
        custo.append([f"{nome}: ganho líquido apurado", brl(dados["realised_gain_brl"])])
        custo.append([f"{nome}: imposto", brl(dados["tax_brl"])])
    if escolhido.get("carried_loss_brl"):
        custo.append(["Prejuízo acumulado informado, abatido na apuração",
                      brl(escolhido["carried_loss_brl"])])
    if not completo:
        custo.append([", ".join(p["ticker"] for p in escolhido["positions_without_cost_basis"])
                      + f": {brl(escolhido['unpriced_sale_brl'])} vendidos sem custo apurável",
                      "não apurado"])
    custo.append(["Total" if completo else "Apurado até aqui (piso)",
                  brl(escolhido["transition_total_brl"])])
    add_table(document, ["Item", "Valor"], custo, widths=[11.6, 5.0])
    nota = ("Apurado por cesta de compensação, ao custo médio: ganhos e prejuízos se encontram "
            "dentro da cesta e nunca entre cestas. A isenção mensal de R$ 20 mil para ações à "
            "vista foi considerada "
            + ("aplicável" if escolhido["exempt_month_assumed"] else "indisponível")
            + " no mês da execução.")
    if "fora_do_escopo" in escolhido["tax_by_bucket"]:
        # Zero na coluna do imposto é a leitura errada mais provável do
        # documento inteiro: não é isenção, é uma conta que não foi feita.
        nota += (" O imposto da cesta fora do escopo aparece como zero porque não é apurado aqui: "
                 "cripto e ativos afins têm regime próprio, e a apuração cabe a quem cuida dela.")
    para(document, nota, size=8.5, color=MUTED)

    if escolhido.get("tax_left_on_table_brl"):
        para(document, f"Este caminho mantém {brl(escolhido['unrealised_loss_kept_brl'])} de "
                       f"prejuízo sem realizar. Realizá-lo abateria "
                       f"{brl(escolhido['tax_left_on_table_brl'])} do imposto acima.",
             size=8.5, color=MUTED)

    heading(document, "5 · O plano que não foi escolhido")
    para(document, "Registrado porque uma decisão sem alternativa documentada não se distingue, "
                   "depois, de uma execução automática.", size=9)
    add_table(document, ["", escolhido["path_label"], outro["path_label"]],
              [["Módulos aplicados", " + ".join(escolhido["modules"]), " + ".join(outro["modules"])],
               ["Custo hoje", brl(escolhido["transition_total_brl"]), brl(outro["transition_total_brl"])],
               ["Volume movimentado", brl(escolhido["turnover_brl"]), brl(outro["turnover_brl"])],
               ["Histórico publicado descreve",
                "sim" if escolhido["track_record_applies"] else "não",
                "sim" if outro["track_record_applies"] else "não"]],
              widths=[4.6, 6.0, 6.0])

    if escolhido.get("fgc_breaches"):
        heading(document, "6 · Risco de crédito acima da cobertura")
        add_table(document, ["Conglomerado", "Posição", "Acima do teto de R$ 250 mil"],
                  [[k, brl(v), brl(v - 250_000)] for k, v in escolhido["fgc_breaches"].items()],
                  widths=[6.0, 5.3, 5.3])
        para(document, "O excedente não tem cobertura do Fundo Garantidor de Créditos. É risco "
                       "de crédito assumido, e assumi-lo precisa ser decisão registrada, não "
                       "consequência de não ter olhado.", size=8.5, color=MUTED)

    heading(document, "7 · O que este documento não afirma")
    limitacoes = [
        "Não estima em quanto tempo a mudança se recupera: exigiria projetar retorno futuro, e a "
        "calibração publicada mostra que projeções assim erram na direção de quem as faz.",
        "Não é ordem de compra ou venda. Nenhuma ordem é transmitida por este sistema.",
        "Os preços são de fechamento e os custos de execução são modelados; a conciliação com a "
        "nota de corretagem acontece depois da operação e pode divergir.",
        "A política que define o destino é retrospectiva na janela em que foi medida; a amostra "
        "confirmatória começa no primeiro pregão de 2027.",
    ]
    if not completo:
        limitacoes.insert(0, "O imposto é parcial e o total é piso: há posição sem custo de "
                             "aquisição apurável, e a diferença não foi estimada.")
    if escolhido.get("issuer_cap_rule"):
        limitacoes.insert(1, f"O teto de concentração deste plano é {pct(escolhido['issuer_cap'])} "
                             f"por emissor: {escolhido['issuer_cap_rule']}.")
    if escolhido.get("locked_tickers"):
        limitacoes.insert(1, "Posições travadas pelo cliente permanecem fora do alcance da "
                             "política: " + ", ".join(escolhido["locked_tickers"]) + ".")
    for item in limitacoes:
        p = para(document, f"•  {item}", size=8.5, color=MUTED, space_after=3)
        p.paragraph_format.left_indent = Cm(0.3)

    para(document, "", size=6)
    heading(document, "8 · Assinatura")
    para(document, "A política calcula e este documento registra. Executar é decisão de quem "
                   "assina, e é a assinatura que torna a decisão defensável depois.", size=9)
    assinatura = document.add_table(rows=2, cols=2)
    assinatura.style = "Table Grid"
    campos = [(("Aprovado por", ""), ("Data", "")),
              (("Cargo e registro", ""), ("Cliente", cliente))]
    for i, linha in enumerate(campos):
        for j, (rotulo, valor) in enumerate(linha):
            cell = assinatura.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(rotulo.upper())
            r.font.size = Pt(7.5); r.bold = True; r.font.color.rgb = MUTED
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            v = p.add_run(valor or " ")
            v.font.size = Pt(10)

    para(document, "", size=6)
    kicker(document, "Verificação")
    linha = (f"Política {registro['policy']} · SHA-256 {registro['registration_sha256']} · "
             f"congelada por {registro['approved_by']} · decisão de destino "
             f"{payload['target_decision']} · mapa por portfolio_mapping.py · perfil por "
             f"client_intake.py.")
    if payload.get("record"):
        # O documento aponta para o registro que o originou. Sem isso a cadeia
        # tem um vão no meio: dá para conferir a política e as contas, e não dá
        # para provar que as respostas eram estas.
        corpo = json.dumps(payload["record"], ensure_ascii=False, sort_keys=True,
                           separators=(",", ":")).encode("utf-8")
        linha += (f" Registro da decisão {payload['record']['schema']} · SHA-256 "
                  f"{hashlib.sha256(corpo).hexdigest()}")
        if payload["record"].get("decided_at"):
            linha += f" · respondido em {payload['record']['decided_at']}"
        linha += "."
    para(document, linha, size=7.5, color=MUTED)
    return document


def payload_from_record(record, positions=None) -> dict:
    """Refaz as contas a partir do registro da decisão.

    Nenhum número atravessa a fronteira entre a tela e o dossiê: o registro traz
    respostas, custos declarados e a escolha, e tudo é recalculado aqui com o
    mesmo módulo de sempre. É o que impede o documento de discordar da tela por
    acidente — não existem dois cálculos para discordarem.
    """
    from client_intake import as_json
    from plan_record import apply_declared_costs
    from portfolio_mapping import adapt_portfolio, map_portfolio
    from research_portfolio_mapping import alvo_do_perfil, carteira_exemplo

    posicoes = apply_declared_costs(list(positions or carteira_exemplo()),
                                    record.declared_costs)
    alvo, decisao = alvo_do_perfil(record.profile)
    return {
        "status": "generated_from_plan_record",
        "target_profile": record.profile,
        "target_decision": decisao,
        "record": record.to_json(),
        "intake": {"answers": record.answers,
                   "assessment": record.intake.assessment(),
                   "questionnaire": as_json()},
        "mapping": map_portfolio(posicoes, alvo),
        "alternative": adapt_portfolio(posicoes, alvo),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--record", type=Path,
                        help="registro da decisão emitido pela tela (plan_record_v1). "
                             "Quando presente, manda no perfil, nos custos e no plano.")
    parser.add_argument("--mapping", type=Path,
                        default=ROOT / "artifacts/portfolio_mapping_v1/mapping_example.json")
    parser.add_argument("--caminho", choices=("adequar", "adaptar", "ambos"), default="ambos")
    parser.add_argument("--cliente", default="Cliente de demonstração")
    parser.add_argument("--out", type=Path, default=OUT_DOCX)
    args = parser.parse_args()

    registro = json.loads((ROOT / "data/benevente_profile_ladder_v3_registration.json")
                          .read_text(encoding="utf-8"))
    args.out.mkdir(parents=True, exist_ok=True)

    if args.record:
        from plan_record import load
        decisao = load(args.record)
        payload = payload_from_record(decisao)
        cliente = decisao.client or args.cliente
        # A escolha é de quem assina, não do argumento de linha de comando.
        caminhos = (decisao.chosen_path,)
    else:
        payload = json.loads(args.mapping.read_text(encoding="utf-8"))
        cliente = args.cliente
        caminhos = ("adequar", "adaptar") if args.caminho == "ambos" else (args.caminho,)

    for escolha in caminhos:
        document = build(payload, cliente, escolha, registro)
        destino = args.out / f"plano_{payload['target_profile']}_{escolha}.docx"
        document.save(destino)
        mapa = payload["mapping"] if escolha == "adequar" else payload["alternative"]
        completo = mapa.get("tax_is_complete", True)
        print(f"{destino.name}: {'' if completo else 'a partir de '}"
              f"{brl(mapa['transition_total_brl'])} ({pct(mapa['transition_cost_pct'], 2)}) · "
              f"{'histórico aplica' if mapa['track_record_applies'] else 'histórico NÃO aplica'}"
              f"{'' if completo else ' · imposto incompleto'}")


if __name__ == "__main__":
    main()
