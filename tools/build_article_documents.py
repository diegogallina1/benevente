"""Build the two reviewed manuscripts from their canonical text sources."""
from __future__ import annotations

import re
from zipfile import ZIP_DEFLATED, ZipFile
from pathlib import Path
from shutil import copy2
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
BTECH_SOURCE = ROOT / "paper" / "fucape_btech_2026.md"
BTECH_TEMPLATE = ROOT / "paper" / "templates" / "BTECH_2026_official.docx"
BTECH_ARCHITECTURE = ROOT / "paper" / "assets" / "btech_architecture.png"
BTECH_OUTPUT = OUTPUTS / "Benevente_Wealth_System_BTECH_Final.docx"
IEEE_SOURCE = ROOT / "paper" / "ieee_cifer_2027.tex"
IEEE_OUTPUT = OUTPUTS / "Benevente_Quant_AI_IEEE_Final.tex"


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def sanitize_core_properties(path: Path) -> None:
    clean_path = path.with_suffix(".clean.docx")
    core_namespace = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
    dc_namespace = "http://purl.org/dc/elements/1.1/"
    with ZipFile(path, "r") as source, ZipFile(clean_path, "w", ZIP_DEFLATED) as target:
        seen: set[str] = set()
        for item in source.infolist():
            if item.filename in seen:
                continue
            seen.add(item.filename)
            data = source.read(item)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                values = {
                    f"{{{dc_namespace}}}creator": "",
                    f"{{{core_namespace}}}lastModifiedBy": "",
                    f"{{{dc_namespace}}}title": "Benevente Wealth System",
                    f"{{{dc_namespace}}}subject": "Manuscrito tecnológico BTech 2026",
                }
                for tag, value in values.items():
                    element = root.find(tag)
                    if element is None:
                        element = ET.SubElement(root, tag)
                    element.text = value
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename in {"word/header1.xml", "word/footer1.xml"}:
                root = ET.fromstring(data)
                description = (
                    "Identidade visual do Business Tech Congress 2026"
                    if item.filename == "word/header1.xml"
                    else "Logotipo da Fucape Business School"
                )
                for element in root.iter():
                    if element.tag.endswith("}docPr"):
                        element.set("descr", description)
                        element.set("title", description)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, data)
    clean_path.replace(path)


TABLE_TITLES = (
    "Camadas do artefato e respectivas saídas auditáveis",
    "Diagnóstico retrospectivo da carteira e dos comparadores entre 2015 e 2025",
    "Exposição anual a séries com proventos imputados",
    "Incerteza do excesso de retorno em reamostragem pareada",
    "Estatísticas de correção por múltiplas tentativas",
    "Desempenho por cadência de reseleção da carteira",
    "Comparações do experimento com modelo de linguagem",
    "Defeitos identificados na auditoria interna e respectivas correções",
    "Matriz de afirmações, evidências e situação",
)


def set_border(properties, edge: str, *, value: str = "single", size: str = "8") -> None:
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders"); properties.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}"); borders.append(border)
    border.set(qn("w:val"), value)
    if value != "nil":
        border.set(qn("w:sz"), size); border.set(qn("w:color"), "000000")


def set_cell_bottom_border(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000"); borders.append(bottom)


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW"); properties.append(width)
    width.set(qn("w:w"), str(width_dxa)); width.set(qn("w:type"), "dxa")


def set_table_geometry(table, rows: list[list[str]], total_width_dxa: int = 9060) -> None:
    column_count = len(rows[0])
    scores = []
    for column in range(column_count):
        longest = max(len(re.sub(r"[*`]", "", row[column])) for row in rows)
        scores.append(max(8, min(longest, 34)))
    score_total = sum(scores)
    widths = [max(900, round(total_width_dxa * score / score_total)) for score in scores]
    difference = total_width_dxa - sum(widths)
    widths[-1] += difference

    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW"); properties.append(table_width)
    table_width.set(qn("w:w"), str(total_width_dxa)); table_width.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths:
        column = OxmlElement("w:gridCol"); column.set(qn("w:w"), str(width_dxa)); grid.append(column)
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths):
            set_cell_width(cell, width_dxa)


def ensure_paragraph_style(document: Document, name: str):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def add_decimal_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    number_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    level = OxmlElement("w:lvl"); level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "decimal"); level.append(fmt)
    text = OxmlElement("w:lvlText"); text.set(qn("w:val"), "%1."); level.append(text)
    justification = OxmlElement("w:lvlJc"); justification.set(qn("w:val"), "left"); level.append(justification)
    properties = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind"); indent.set(qn("w:left"), "720"); indent.set(qn("w:hanging"), "360")
    properties.append(indent); level.append(properties); abstract.append(level); numbering.append(abstract)

    number = OxmlElement("w:num"); number.set(qn("w:numId"), str(number_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref); numbering.append(number)
    return number_id


def apply_numbering(paragraph, number_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl"); level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId"); identifier.set(qn("w:val"), str(number_id))
    number_properties.append(level); number_properties.append(identifier); properties.append(number_properties)


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            # Reserve bold for structural hierarchy such as headings and table
            # headers. Inline emphasis is rendered as normal body text.
            paragraph.add_run(token[2:-2])
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1]); run.font.name = "Courier New"; run.font.size = Pt(9)
        else:
            run = paragraph.add_run(token[1:-1]); run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_table(document: Document, lines: list[str], table_number: int) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for index, row in enumerate(rows) if index != 1]

    number = document.add_paragraph()
    number.paragraph_format.keep_with_next = True
    number.paragraph_format.space_after = Pt(0)
    run = number.add_run(f"Tabela {table_number}"); run.bold = True
    title = document.add_paragraph()
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(TABLE_TITLES[table_number - 1]); run.italic = True

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, rows)
    table_properties = table._tbl.tblPr
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(table_properties, edge, value="nil")
    set_border(table_properties, "top"); set_border(table_properties, "bottom")
    for row_index, values in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit"); row_properties.append(cannot_split)
        if row_index == 0:
            repeat_header = OxmlElement("w:tblHeader"); repeat_header.set(qn("w:val"), "true"); row_properties.append(repeat_header)
        for cell, value in zip(table.rows[row_index].cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = row_index < len(rows) - 1
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(8.5)
                if row_index == 0: run.bold = True
            if row_index == 0: set_cell_bottom_border(cell)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(3); note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.0
    label = note.add_run("Nota. "); label.italic = True
    run = note.add_run("Elaboração própria com base nos artefatos reproduzíveis da pesquisa.")
    for item in (label, run):
        item.font.name = "Times New Roman"; item.font.size = Pt(8)


def add_figure(document: Document, figure_number: int) -> None:
    if not BTECH_ARCHITECTURE.exists():
        raise FileNotFoundError(f"Missing architecture figure: {BTECH_ARCHITECTURE}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    shape = paragraph.add_run().add_picture(str(BTECH_ARCHITECTURE), width=Cm(15.5))
    description = "Fluxo do Benevente: dados B3 e CVM passam por validação, seleção quantitativa, alocação e revisão humana; o modelo de linguagem apenas explica fatos aprovados; o dossiê reúne todo o registro."
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", "Arquitetura auditável do Benevente Wealth System")

    number = document.add_paragraph()
    number.paragraph_format.first_line_indent = Cm(0)
    number.paragraph_format.keep_with_next = True
    number.paragraph_format.space_after = Pt(0)
    run = number.add_run(f"Figura {figure_number}"); run.bold = True
    title = document.add_paragraph()
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Arquitetura e separação de responsabilidades do artefato"); run.italic = True
    note = document.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.line_spacing = 1.0
    note.paragraph_format.space_after = Pt(8)
    label = note.add_run("Nota. "); label.italic = True
    note.add_run("As setas indicam o fluxo de dados. Não existe caminho do modelo de linguagem para alterar pesos ou aprovar ordens.")
    for item in (number, title, note):
        for item_run in item.runs:
            item_run.font.name = "Times New Roman"
            item_run.font.size = Pt(8)


def markdown_blocks(text: str):
    lines = text.splitlines(); index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line or line.strip() == "---": index += 1; continue
        if line == "[[FIGURE:architecture]]": yield "figure", "architecture"; index += 1; continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            table = [line, lines[index + 1]]; index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table.append(lines[index]); index += 1
            yield "table", table; continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#")); yield f"h{level}", line[level:].strip(); index += 1; continue
        if re.match(r"^[-*] ", line): yield "bullet", re.sub(r"^[-*] ", "", line); index += 1; continue
        if re.match(r"^\d+\. ", line): yield "number", line; index += 1; continue
        paragraph = [line]; index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if (not next_line or next_line.startswith("#") or next_line.startswith("|") or
                    re.match(r"^[-*] ", next_line) or re.match(r"^\d+\. ", next_line) or next_line.strip() == "---"):
                break
            paragraph.append(next_line); index += 1
        yield "paragraph", " ".join(paragraph)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, bold in (("Heading 1", True), ("Heading 2", False), ("Heading 3", False)):
        style = ensure_paragraph_style(document, name)
        style.font.name = "Times New Roman"; style.font.size = Pt(12); style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(9); style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.keep_with_next = True


def build_btech() -> None:
    if not BTECH_TEMPLATE.exists(): raise FileNotFoundError("The retained BTech-formatted DOCX is missing.")
    copy2(BTECH_TEMPLATE, BTECH_OUTPUT)
    document = Document(BTECH_OUTPUT); clear_body(document); configure_document(document)
    in_references = False
    in_summary = False
    body_started = False
    decimal_number_id: int | None = None
    previous_kind: str | None = None
    table_number = 0
    figure_number = 0
    for kind, content in markdown_blocks(BTECH_SOURCE.read_text(encoding="utf-8")):
        if kind == "table":
            table_number += 1
            add_table(document, content, table_number)
            previous_kind = kind
            continue
        if kind == "figure":
            figure_number += 1
            add_figure(document, figure_number)
            previous_kind = kind
            continue
        if kind == "h1":
            paragraph = document.add_paragraph(style="Normal"); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.line_spacing = 1.0
            content = content.upper()
        elif kind == "h2" and content == "Resumo":
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(6); paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_with_next = True
            in_summary = True
        elif kind in {"h2", "h3"}:
            if kind == "h2" and not body_started:
                document.add_page_break(); body_started = True; in_summary = False
            paragraph = document.add_paragraph(style="Heading 1" if kind == "h2" else "Heading 2")
            content = content.upper()
            if kind == "h2" and content == "REFERÊNCIAS":
                in_references = True
        elif kind in {"bullet", "number"}:
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            if kind == "number":
                if previous_kind != "number": decimal_number_id = add_decimal_numbering(document)
                apply_numbering(paragraph, decimal_number_id)
                content = re.sub(r"^\d+\.\s+", "", content)
            else:
                apply_numbering(paragraph, 1)
        else:
            paragraph = document.add_paragraph(style="Normal")
            if in_summary:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(6)
                if content.startswith("Palavras-chave:"):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if in_references:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Cm(0.75)
                paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        add_inline(paragraph, content)
        if kind == "h1":
            for run in paragraph.runs: run.bold = True
        elif kind == "h2" and content == "Resumo":
            for run in paragraph.runs: run.bold = True
        previous_kind = kind
    if table_number != len(TABLE_TITLES):
        raise ValueError(f"Expected {len(TABLE_TITLES)} BTech tables, found {table_number}")
    if figure_number != 1:
        raise ValueError(f"Expected 1 BTech figure, found {figure_number}")
    temporary = BTECH_OUTPUT.with_suffix(".tmp.docx")
    document.save(temporary); sanitize_core_properties(temporary); temporary.replace(BTECH_OUTPUT)


def build_ieee() -> None:
    text = IEEE_SOURCE.read_text(encoding="utf-8")
    if "\\documentclass[10pt,conference]{IEEEtran}" not in text: raise ValueError("IEEE manuscript is not using the conference template")
    if "MVO is an independent benchmark" not in text: raise ValueError("The canonical role contract is missing from the IEEE manuscript")
    copy2(IEEE_SOURCE, IEEE_OUTPUT)


if __name__ == "__main__":
    OUTPUTS.mkdir(exist_ok=True); build_btech(); build_ieee()
    print(BTECH_OUTPUT); print(IEEE_OUTPUT)
